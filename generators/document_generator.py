"""
Document Generator Module
Handles Word, PDF, and Text document generation
"""

import io
import tempfile
from docx import Document
from docx.shared import Inches

# Try to import reportlab, with fallback
try:
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False
    print("⚠️  Warning: ReportLab not installed. PDF generation will be disabled.")


class DocumentGenerator:
    def __init__(self, data_generator):
        self.data_generator = data_generator

    def generate_document_content(self, content_type, pages, subject):
        """Generate document content - simple approach for short documents"""
        target_words = pages * 275
        
        prompts = {
            "whitepaper": f"""Write a comprehensive technical whitepaper on {subject} that is EXACTLY {pages} pages long (approximately {target_words} words).

STRUCTURE REQUIREMENTS:
- Page 1: Executive Summary (1 full page)
- Page 2: Introduction and Background (1 full page)
{f"- Pages 3-{pages-2}: Technical Analysis, Methodology, Implementation Details ({pages-4} pages)" if pages > 4 else ""}
{f"- Page {pages-1}: Findings and Results (1 full page)" if pages > 3 else ""}
- Page {pages}: Conclusions and References (1 full page)

CONTENT REQUIREMENTS:
- Write detailed paragraphs with 4-6 sentences each
- Include specific technical details and examples
- Add subsections with descriptive headings
- Ensure each section is substantive and detailed
- The document must be exactly {pages} pages when printed
- Target approximately {target_words} total words

Make it detailed, professional, and comprehensive. Note at the end that this is synthetically created data.""",

            "article": f"""Write a detailed article about {subject} that is EXACTLY {pages} pages long (approximately {target_words} words).

STRUCTURE REQUIREMENTS:
- Page 1: Introduction and overview
{f"- Pages 2-{pages-1}: Main content with multiple detailed sections ({pages-2} pages)" if pages > 2 else ""}
- Page {pages}: Conclusion and key takeaways

CONTENT REQUIREMENTS:
- Write in-depth paragraphs with 5-7 sentences each
- Include practical examples and case studies
- Add multiple subsections with detailed explanations
- Use bullet points and numbered lists where appropriate
- The article must be exactly {pages} pages when printed
- Target approximately {target_words} total words

Include multiple sections, subsections, and practical examples. Note at the end that this is synthetically created data.""",

            "report": f"""Write a comprehensive business report on {subject} that is EXACTLY {pages} pages long (approximately {target_words} words).

STRUCTURE REQUIREMENTS:
- Page 1: Executive Summary and Key Findings
- Page 2: Introduction and Methodology
{f"- Pages 3-{pages-2}: Detailed Analysis and Data ({pages-4} pages)" if pages > 4 else ""}
{f"- Page {pages-1}: Strategic Recommendations (1 full page)" if pages > 3 else ""}
- Page {pages}: Conclusions and Next Steps

CONTENT REQUIREMENTS:
- Include detailed data analysis descriptions
- Add charts and graphs descriptions (describe what they would show)
- Write comprehensive strategic recommendations
- Include market analysis and competitive landscape
- The report must be exactly {pages} pages when printed
- Target approximately {target_words} total words

Make it detailed and professional with strategic recommendations. Note at the end that this is synthetically created data.""",

            "proposal": f"""Write a detailed project proposal for {subject} that is EXACTLY {pages} pages long (approximately {target_words} words).

STRUCTURE REQUIREMENTS:
- Page 1: Project Overview and Objectives
- Page 2: Detailed Scope and Requirements
{f"- Page 3: Timeline and Milestones" if pages > 3 else ""}
{f"- Page 4: Budget and Resource Allocation" if pages > 4 else ""}
{f"- Pages {5 if pages > 4 else 3}-{pages-1}: Implementation Plan and Risk Analysis" if pages > 4 else ""}
- Page {pages}: Expected Outcomes and Success Metrics

CONTENT REQUIREMENTS:
- Detailed project scope with specific deliverables
- Comprehensive timeline with multiple phases
- Detailed budget breakdown with justifications
- Risk analysis with mitigation strategies
- The proposal must be exactly {pages} pages when printed
- Target approximately {target_words} total words

Include scope, timeline, budget considerations and risk analysis. Note at the end that this is synthetically created data.""",

            "design": f"""Write a detailed design document for an IT project on {subject} that is EXACTLY {pages} pages long (approximately {target_words} words).

STRUCTURE REQUIREMENTS:
- Page 1: System Overview and Architecture
- Page 2: Technical Requirements and Components
{f"- Page 3: Interface Design and Data Flow" if pages > 3 else ""}
{f"- Page 4: Security and Performance Considerations" if pages > 4 else ""}
{f"- Pages {5 if pages > 4 else 3}-{pages-1}: Implementation Details and Testing" if pages > 4 else ""}
- Page {pages}: Deployment and Maintenance

CONTENT REQUIREMENTS:
- Detailed system architecture descriptions
- Technical component specifications
- Interface design with specific examples
- Data flow diagrams descriptions
- Security protocols and performance metrics
- The document must be exactly {pages} pages when printed
- Target approximately {target_words} total words

Include architecture, components, interfaces and data flow. Note at the end that this is synthetically created data."""
        }

        if content_type not in prompts:
            raise ValueError(f"Invalid content type: {content_type}. Choose from {list(prompts.keys())}.")

        prompt = prompts.get(content_type)
        return self.data_generator.generate_with_ollama(prompt, max_tokens=6000)

    def generate_document_content_iterative(self, content_type, pages, subject):
        """Generate document content by building it section by section for longer documents"""
        
        print(f"🔄 Generating {pages}-page {content_type} iteratively...")
        sections = []
        
        # Section configurations for different document types
        section_configs = {
            "whitepaper": [
                ("Executive Summary", "Write a comprehensive 1-page executive summary with key findings and recommendations"),
                ("Introduction and Background", "Write a detailed introduction covering the background, context, and importance of the topic"),
                ("Technical Analysis", "Write an in-depth technical analysis section with detailed explanations and technical specifications"),
                ("Methodology", "Write a detailed methodology section explaining approaches, frameworks, and implementation strategies"),
                ("Implementation Details", "Write comprehensive implementation details including step-by-step processes and technical requirements"),
                ("Results and Findings", "Write detailed results and findings with data analysis and performance metrics"),
                ("Future Considerations", "Write about future implications, scalability, and evolution of the technology"),
                ("Conclusions and References", "Write comprehensive conclusions with actionable recommendations and references")
            ],
            "article": [
                ("Introduction", "Write a comprehensive introduction that sets the context and engages the reader"),
                ("Main Analysis", "Write the main analysis section with detailed content and thorough examination"),
                ("Case Studies and Examples", "Write detailed case studies and real-world examples with specific scenarios"),
                ("Industry Impact and Implications", "Write about industry impact, market implications, and economic effects"),
                ("Current Trends and Developments", "Write about current trends, recent developments, and emerging patterns"),
                ("Best Practices and Recommendations", "Write about best practices, recommendations, and actionable insights"),
                ("Future Outlook", "Write about future trends, predictions, and long-term implications"),
                ("Conclusion", "Write a comprehensive conclusion that summarizes key points and provides final thoughts")
            ],
            "report": [
                ("Executive Summary", "Write an executive summary with key findings, recommendations, and critical insights"),
                ("Introduction and Methodology", "Write introduction covering scope, objectives, and detailed methodology"),
                ("Market Analysis", "Write detailed market analysis including size, trends, competitors, and opportunities"),
                ("Data Analysis and Insights", "Write comprehensive data analysis with statistical insights and interpretations"),
                ("Strategic Recommendations", "Write strategic recommendations with detailed implementation guidance"),
                ("Risk Assessment", "Write thorough risk assessment including potential challenges and mitigation strategies"),
                ("Implementation Plan", "Write detailed implementation plan with timelines, resources, and success metrics"),
                ("Financial Analysis", "Write financial analysis including costs, benefits, and ROI projections"),
                ("Conclusions and Next Steps", "Write conclusions with clear next steps and action items")
            ],
            "proposal": [
                ("Project Overview and Objectives", "Write project overview covering goals, scope, and strategic alignment"),
                ("Scope and Requirements", "Write detailed scope definition and comprehensive requirements analysis"),
                ("Timeline and Milestones", "Write comprehensive timeline with detailed milestones and deliverable schedules"),
                ("Budget and Resource Allocation", "Write detailed budget breakdown and resource allocation plans"),
                ("Implementation Strategy", "Write implementation strategy with detailed methodology and approach"),
                ("Risk Analysis and Mitigation", "Write comprehensive risk analysis with detailed mitigation strategies"),
                ("Team and Expertise", "Write about team composition, expertise, and organizational capabilities"),
                ("Quality Assurance", "Write about quality assurance processes, testing, and validation procedures"),
                ("Expected Outcomes and Success Metrics", "Write about expected outcomes, success criteria, and measurement methods")
            ],
            "design": [
                ("System Overview and Architecture", "Write system overview covering high-level architecture and design principles"),
                ("Technical Requirements", "Write detailed technical requirements including functional and non-functional specifications"),
                ("Interface Design and User Experience", "Write about interface design, user experience, and interaction patterns"),
                ("Data Architecture and Flow", "Write about data architecture, database design, and information flow"),
                ("Security and Performance", "Write about security considerations, performance requirements, and scalability"),
                ("Implementation Details", "Write detailed implementation specifications including technologies and frameworks"),
                ("Testing and Quality Assurance", "Write about testing strategies, quality assurance processes, and validation methods"),
                ("Deployment and Infrastructure", "Write about deployment architecture, infrastructure requirements, and operational considerations"),
                ("Maintenance and Support", "Write about maintenance procedures, support processes, and long-term sustainability")
            ]
        }
        
        # Get sections for this document type
        all_sections = section_configs.get(content_type, section_configs["article"])
        
        # Calculate sections to use based on page count
        sections_to_use = min(len(all_sections), max(pages, 3))
        selected_sections = all_sections[:sections_to_use]
        
        # If we need more sections for very long documents, add additional analysis sections
        while len(selected_sections) < pages:
            additional_sections = [
                ("Additional Technical Analysis", "Write additional detailed technical analysis with deeper insights"),
                ("Supplementary Research", "Write supplementary research findings and supporting evidence"),
                ("Extended Case Studies", "Write extended case studies with more detailed examples"),
                ("Advanced Considerations", "Write about advanced considerations and complex scenarios")
            ]
            selected_sections.extend(additional_sections[:pages - len(selected_sections)])
        
        # Calculate words per section (aim for more words to ensure length)
        total_target_words = pages * 300  # Slightly higher target
        words_per_section = total_target_words // len(selected_sections)
        
        # Generate each section
        for i, (section_title, section_instruction) in enumerate(selected_sections):
            section_prompt = f"""{section_instruction} for a {content_type} about {subject}.

REQUIREMENTS:
- Write approximately {words_per_section} words for this section
- Include detailed explanations with specific examples related to {subject}
- Use professional, technical language appropriate for the topic
- Make this section substantial and comprehensive with multiple paragraphs
- Include subsections, bullet points, and numbered lists where appropriate
- Provide concrete examples and detailed analysis
- This is section {i+1} of {len(selected_sections)} in a {pages}-page document

Topic: {subject}
Document Type: {content_type}
Section: {section_title}

Write detailed, professional content that thoroughly covers this section with substantial depth and analysis."""

            print(f"📝 Generating section {i+1}/{len(selected_sections)}: {section_title}")
            section_content = self.data_generator.generate_with_ollama(section_prompt, max_tokens=2000)
            
            # Add section to document with proper formatting
            formatted_section = f"# {section_title}\n\n{section_content}"
            sections.append(formatted_section)
        
        # Add final disclaimer
        final_note = "\n\n# Disclaimer\n\nThis document has been synthetically generated using AI for demonstration purposes. All content, data, recommendations, and analysis are artificially created and should not be used for actual business decisions, implementation, or as factual reference material. Please consult appropriate experts and conduct proper research for real-world applications."
        sections.append(final_note)
        
        # Combine all sections
        full_content = "\n\n".join(sections)
        
        word_count = len(full_content.split())
        print(f"✅ Generated document with {len(sections)-1} sections, approximately {word_count} words")
        return full_content

    def create_word_document(self, content):
        """Create a Word document with the generated content"""
        doc = Document()
        
        # Split content into paragraphs and add to document
        paragraphs = content.split('\n\n')
        
        for i, paragraph in enumerate(paragraphs):
            if paragraph.strip():
                if i == 0:  # First paragraph as title
                    doc.add_heading(paragraph.strip()[:100], 0)
                elif paragraph.strip().startswith('# '):
                    # Main headings
                    heading_text = paragraph.strip().replace('# ', '')
                    doc.add_heading(heading_text, 1)
                elif paragraph.strip().startswith('## '):
                    # Sub headings
                    heading_text = paragraph.strip().replace('## ', '')
                    doc.add_heading(heading_text, 2)
                elif paragraph.strip().startswith('- ') or paragraph.strip().startswith('* '):
                    # Bullet points
                    doc.add_paragraph(paragraph.strip(), style='List Bullet')
                elif paragraph.strip().startswith(('1. ', '2. ', '3. ', '4. ', '5. ')):
                    # Numbered lists
                    doc.add_paragraph(paragraph.strip(), style='List Number')
                else:
                    # Regular paragraphs
                    doc.add_paragraph(paragraph.strip())
        
        doc_bytes = io.BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)
        return doc_bytes.getvalue()

    def create_pdf_document(self, content):
        """Create a PDF document with the generated content"""
        if not REPORTLAB_AVAILABLE:
            raise ImportError("ReportLab library is not installed. Please install it with: pip install reportlab")
        
        pdf_bytes = io.BytesIO()
        doc = SimpleDocTemplate(pdf_bytes, pagesize=letter)
        
        # Get styles
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            spaceAfter=30,
        )
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=14,
            spaceAfter=12,
        )
        subheading_style = ParagraphStyle(
            'CustomSubheading',
            parent=styles['Heading3'],
            fontSize=12,
            spaceAfter=8,
        )
        
        story = []
        paragraphs = content.split('\n\n')
        
        for i, paragraph in enumerate(paragraphs):
            if paragraph.strip():
                if i == 0:  # First paragraph as title
                    story.append(Paragraph(paragraph.strip()[:100], title_style))
                    story.append(Spacer(1, 12))
                elif paragraph.strip().startswith('# '):
                    # Main headings
                    heading_text = paragraph.strip().replace('# ', '')
                    story.append(Paragraph(heading_text, heading_style))
                    story.append(Spacer(1, 6))
                elif paragraph.strip().startswith('## '):
                    # Sub headings
                    heading_text = paragraph.strip().replace('## ', '')
                    story.append(Paragraph(heading_text, subheading_style))
                    story.append(Spacer(1, 4))
                else:
                    # Regular paragraphs
                    story.append(Paragraph(paragraph.strip(), styles['Normal']))
                    story.append(Spacer(1, 12))
        
        doc.build(story)
        pdf_bytes.seek(0)
        return pdf_bytes.getvalue()

    def create_text_document(self, content):
        """Create a plain text document with the generated content"""
        # Clean up the content for plain text
        text_content = content.replace('# ', '')
        text_content = text_content.replace('## ', '')
        
        # Add a simple header
        header = "="*80 + "\n"
        header += "SYNTHETIC DOCUMENT - GENERATED CONTENT\n" 
        header += "="*80 + "\n\n"
        
        final_content = header + text_content
        return final_content.encode('utf-8')

    def generate_document(self, content_type, pages, subject, file_format):
        """Main document generation orchestrator"""
        # Use iterative approach for longer documents (3+ pages)
        if pages >= 3:
            content = self.generate_document_content_iterative(content_type, pages, subject)
            print(f"✅ Used iterative approach for {pages} pages")
        else:
            content = self.generate_document_content(content_type, pages, subject)
            print(f"✅ Used standard approach for {pages} pages")
        
        if file_format == "Word Document (.docx)":
            file_bytes = self.create_word_document(content)
            suffix = '.docx'
        elif file_format == "PDF Document (.pdf)":
            file_bytes = self.create_pdf_document(content)
            suffix = '.pdf'
        else:  # Text File (.txt)
            file_bytes = self.create_text_document(content)
            suffix = '.txt'
        
        # Save to temporary file and return path
        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix, prefix='synthetic_') as tmp_file:
            tmp_file.write(file_bytes)
            temp_path = tmp_file.name
        
        return temp_path, f"✅ Generated {pages}-page {content_type} about '{subject}' successfully!"
